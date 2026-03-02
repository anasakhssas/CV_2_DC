"""Génération du Dossier de Compétences au format Word (.docx).

Mise en page fidèle au template ACHMITECH :
  Page 1 — Photo · Titre · Infos candidat · Top5 Hard/Soft/Outils/Langues
  Page 2+ — Détail des formations et expériences professionnelles
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

from app.models import DossierCompetences
from app.services.years_calculator import calculate_years_of_experience

logger = logging.getLogger("cv2dc.docx_generator")

# ── Palette couleurs ACHMITECH ───────────────────────────────
TEAL         = "009999"                       # fond labels / entêtes sections
TEAL_RGB     = RGBColor(0x00, 0x99, 0x99)
WHITE        = "FFFFFF"
WHITE_RGB    = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_ROW    = "F2F2F2"                       # lignes paires tableaux
TITLE_RGB    = RGBColor(0x00, 0x80, 0x80)     # "DOSSIER DE COMPETENCES"
GREY_RGB     = RGBColor(0x80, 0x80, 0x80)     # texte secondaire
BLACK_RGB    = RGBColor(0x00, 0x00, 0x00)
BORDER_COLOR = "AAAAAA"

PAGE_W = 17.0   # largeur utile (cm) — A4 avec marges 2 cm



# ─────────────────────────────────────────────────────────────
#  Helpers XML bas niveau
# ─────────────────────────────────────────────────────────────

def _cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_borders(cell, color: str = BORDER_COLOR,
                  sides=("top", "left", "bottom", "right")) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _no_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _table_width(table, width_cm: float) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(width_cm * 567)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _col_width(cell, width_cm: float) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _run(para, text: str, bold: bool = False, italic: bool = False,
         size: int = 10, color: RGBColor = BLACK_RGB,
         font: str = "Calibri") -> None:
    r = para.add_run(text)
    r.font.name   = font
    r.font.size   = Pt(size)
    r.font.bold   = bold
    r.font.italic = italic
    r.font.color.rgb = color


def _stars(level: int, max_l: int = 5) -> str:
    return "★" * int(level) + "☆" * (max_l - int(level))


def _years_label(dossier: DossierCompetences) -> str:
    if not dossier.experiences:
        return "—"
    try:
        yoe   = calculate_years_of_experience(dossier.experiences)
        total = yoe.total_years
        return f"{total} an{'s' if total > 1 else ''}" if total else "—"
    except Exception:
        return "—"


def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._r.append(br)


# ─────────────────────────────────────────────────────────────
#  Bloc Top-5 ACHMITECH
# ─────────────────────────────────────────────────────────────

LEVEL_COL = 3.2
MAIN_COL  = PAGE_W - LEVEL_COL


def _top5_table(doc: Document, title: str,
                rows_data: list[tuple[str, str]]) -> None:
    """
    â”Œâ”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”¬â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”
    │  TITRE                               │ Votre Niveau│
    │                                      │ de 1-5 sur 5│
    ├──────────────────────────────────────┼─────────────┤
    │  skill name                          │  ★★★☆☆      │
    └──────────────────────────────────────┴─────────────┘
    Affiche exactement 5 lignes (complétées avec des vides).
    """
    tbl = doc.add_table(rows=1, cols=2)
    _no_borders(tbl)
    _table_width(tbl, PAGE_W)

    # Entête
    hdr = tbl.rows[0]
    lh, rh = hdr.cells[0], hdr.cells[1]
    _col_width(lh, MAIN_COL)
    _col_width(rh, LEVEL_COL)
    for c in (lh, rh):
        _cell_bg(c, TEAL)
        _cell_borders(c, "007777")

    p_l = lh.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_l.paragraph_format.space_before = Pt(4)
    p_l.paragraph_format.space_after  = Pt(4)
    _run(p_l, title, bold=True, size=10, color=WHITE_RGB)

    p_r = rh.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r.paragraph_format.space_before = Pt(2)
    p_r.paragraph_format.space_after  = Pt(2)
    _run(p_r, "Votre Niveau\nde 1-5 sur 5", bold=True, size=9, color=WHITE_RGB)

    # 5 lignes de données
    padded = (list(rows_data) + [("", "")] * 5)[:5]
    for i, (name_val, level_val) in enumerate(padded):
        row = tbl.add_row()
        bg  = LIGHT_ROW if i % 2 == 0 else WHITE
        nc, lc = row.cells[0], row.cells[1]
        _col_width(nc, MAIN_COL)
        _col_width(lc, LEVEL_COL)
        for c in (nc, lc):
            _cell_bg(c, bg)
            _cell_borders(c, BORDER_COLOR)

        p_n = nc.paragraphs[0]
        p_n.paragraph_format.space_before = Pt(2)
        p_n.paragraph_format.space_after  = Pt(2)
        p_n.paragraph_format.left_indent  = Cm(0.2)
        if name_val:
            _run(p_n, name_val, size=9, color=BLACK_RGB)

        p_lv = lc.paragraphs[0]
        p_lv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_lv.paragraph_format.space_before = Pt(2)
        p_lv.paragraph_format.space_after  = Pt(2)
        if level_val:
            _run(p_lv, level_val, size=9, color=BLACK_RGB)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ─────────────────────────────────────────────────────────────
#  Banner de section (pages 2+)
# ─────────────────────────────────────────────────────────────

def _section_banner(doc: Document, title: str) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    _table_width(tbl, PAGE_W)
    c = tbl.rows[0].cells[0]
    _cell_bg(c, TEAL)
    _cell_borders(c, "007777")
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    _run(p, title, bold=True, size=11, color=WHITE_RGB)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _table_full_borders(table, color: str = "007777") -> None:
    """Outer + inner borders on all sides."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _row_min_height(row, height_cm: float) -> None:
    """Sets a minimum row height (in cm)."""
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_cm * 567)))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def _add_experience_block(doc: Document, exp) -> None:
    """
    Adds one experience card matching the ACHMITECH template:

    ┌─────────────┬────────────────────────────────┬─────────────┐  teal
    │ DATE DEBUT :│ POSTE OCCUPE :                  │ DATE FIN:   │
    │ <date>      │ <position>  ENTREPRISE: <co>    │ <date>      │
    ├─────────────┴────────────────────────────────┴─────────────┤  teal
    │ RÉSUMÉ MISSION :                                            │
    ├─────────────────────────────────────────────────────────────┤  white
    │ <mission_summary>                                           │
    ├─────────────────────────────────────────────────────────────┤  teal
    │ RÉALISATIONS :                                              │
    ├─────────────────────────────────────────────────────────────┤  white
    │ • achievement 1  • achievement 2 …                          │
    ├─────────────────────────────────────────────────────────────┤  teal
    │ ENVIRONNEMENT TECHNIQUE ET TAILLE DE L'EQUIPE :             │
    ├─────────────────────────────────────────────────────────────┤  white
    │ Technologies : …   Méthodologies : …   Taille : …           │
    └─────────────────────────────────────────────────────────────┘
    """
    position  = exp.position   or "Poste non précisé"
    company   = exp.company    or "—"
    date_from = exp.start_date or "?"
    date_to   = exp.end_date   or "Présent"

    DATE_W = 3.2
    MID_W  = PAGE_W - DATE_W * 2   # ≈ 10.6 cm

    # ── Table with 3 cols, 7 rows ──────────────────────────────────
    t = doc.add_table(rows=7, cols=3)
    _table_width(t, PAGE_W)
    _table_full_borders(t, "007777")

    # ── Row 0 : teal 3-column header ──────────────────────────────
    r0 = t.rows[0]
    for ci, cw in enumerate([DATE_W, MID_W, DATE_W]):
        _col_width(r0.cells[ci], cw)
        _cell_bg(r0.cells[ci], TEAL)

    # Left : DATE DEBUT
    pl1 = r0.cells[0].paragraphs[0]
    pl1.paragraph_format.space_before = Pt(3)
    pl1.paragraph_format.space_after  = Pt(1)
    pl1.paragraph_format.left_indent  = Cm(0.2)
    _run(pl1, "DATE DEBUT :", bold=True, size=9, color=WHITE_RGB)
    pl2 = r0.cells[0].add_paragraph()
    pl2.paragraph_format.space_before = Pt(1)
    pl2.paragraph_format.space_after  = Pt(3)
    pl2.paragraph_format.left_indent  = Cm(0.2)
    _run(pl2, date_from, size=9, color=WHITE_RGB)

    # Middle : POSTE OCCUPE / ENTREPRISE
    pm1 = r0.cells[1].paragraphs[0]
    pm1.paragraph_format.space_before = Pt(3)
    pm1.paragraph_format.space_after  = Pt(1)
    pm1.paragraph_format.left_indent  = Cm(0.2)
    _run(pm1, "POSTE OCCUPE :", bold=True, size=9, color=WHITE_RGB)
    pm2 = r0.cells[1].add_paragraph()
    pm2.paragraph_format.space_before = Pt(1)
    pm2.paragraph_format.space_after  = Pt(1)
    pm2.paragraph_format.left_indent  = Cm(0.2)
    _run(pm2, position, size=9, color=WHITE_RGB)
    pm3 = r0.cells[1].add_paragraph()
    pm3.paragraph_format.space_before = Pt(1)
    pm3.paragraph_format.space_after  = Pt(1)
    pm3.paragraph_format.left_indent  = Cm(0.2)
    _run(pm3, "ENTREPRISE :", bold=True, size=9, color=WHITE_RGB)
    pm4 = r0.cells[1].add_paragraph()
    pm4.paragraph_format.space_before = Pt(1)
    pm4.paragraph_format.space_after  = Pt(3)
    pm4.paragraph_format.left_indent  = Cm(0.2)
    _run(pm4, company, size=9, color=WHITE_RGB)

    # Right : DATE FIN
    pr1 = r0.cells[2].paragraphs[0]
    pr1.paragraph_format.space_before = Pt(3)
    pr1.paragraph_format.space_after  = Pt(1)
    pr1.paragraph_format.left_indent  = Cm(0.2)
    _run(pr1, "DATE FIN:", bold=True, size=9, color=WHITE_RGB)
    pr2 = r0.cells[2].add_paragraph()
    pr2.paragraph_format.space_before = Pt(1)
    pr2.paragraph_format.space_after  = Pt(3)
    pr2.paragraph_format.left_indent  = Cm(0.2)
    _run(pr2, date_to, size=9, color=WHITE_RGB)

    # ── Rows 1-6 : full-width merged rows ─────────────────────────

    def _merged_row(row_idx: int, bg: str) -> object:
        """Merge row's 3 cells into one and apply background."""
        row  = t.rows[row_idx]
        cell = row.cells[0].merge(row.cells[2])
        _col_width(cell, PAGE_W)
        _cell_bg(cell, bg)
        return row, cell

    # Row 1 : RÉSUMÉ MISSION label (teal)
    _, cell_rm_lbl = _merged_row(1, TEAL)
    p_rm_lbl = cell_rm_lbl.paragraphs[0]
    p_rm_lbl.paragraph_format.space_before = Pt(3)
    p_rm_lbl.paragraph_format.space_after  = Pt(3)
    p_rm_lbl.paragraph_format.left_indent  = Cm(0.3)
    _run(p_rm_lbl, "RÉSUMÉ MISSION :", bold=True, size=9, color=WHITE_RGB)

    # Row 2 : mission summary content (white)
    row2, cell_rm = _merged_row(2, WHITE)
    _row_min_height(row2, 1.5)
    p_rm = cell_rm.paragraphs[0]
    p_rm.paragraph_format.space_before = Pt(3)
    p_rm.paragraph_format.space_after  = Pt(3)
    p_rm.paragraph_format.left_indent  = Cm(0.3)
    if exp.mission_summary:
        _run(p_rm, exp.mission_summary, size=9, color=BLACK_RGB)

    # Row 3 : RÉALISATIONS label (teal)
    _, cell_re_lbl = _merged_row(3, TEAL)
    p_re_lbl = cell_re_lbl.paragraphs[0]
    p_re_lbl.paragraph_format.space_before = Pt(3)
    p_re_lbl.paragraph_format.space_after  = Pt(3)
    p_re_lbl.paragraph_format.left_indent  = Cm(0.3)
    _run(p_re_lbl, "RÉALISATIONS :", bold=True, size=9, color=WHITE_RGB)

    # Row 4 : achievements content (white)
    row4, cell_re = _merged_row(4, WHITE)
    _row_min_height(row4, 2.5)
    p_re = cell_re.paragraphs[0]
    p_re.paragraph_format.space_before = Pt(3)
    p_re.paragraph_format.space_after  = Pt(1)
    p_re.paragraph_format.left_indent  = Cm(0.3)
    if exp.achievements:
        for i, ach in enumerate(exp.achievements):
            pa = p_re if i == 0 else cell_re.add_paragraph()
            pa.paragraph_format.space_before = Pt(1)
            pa.paragraph_format.space_after  = Pt(1)
            pa.paragraph_format.left_indent  = Cm(0.3)
            _run(pa, f"• {ach}", size=9, color=BLACK_RGB)

    # Row 5 : ENVIRONNEMENT TECHNIQUE label (teal)
    _, cell_et_lbl = _merged_row(5, TEAL)
    p_et_lbl = cell_et_lbl.paragraphs[0]
    p_et_lbl.paragraph_format.space_before = Pt(3)
    p_et_lbl.paragraph_format.space_after  = Pt(3)
    p_et_lbl.paragraph_format.left_indent  = Cm(0.3)
    _run(p_et_lbl, "ENVIRONNEMENT TECHNIQUE ET TAILLE DE L'EQUIPE :",
         bold=True, size=9, color=WHITE_RGB)

    # Row 6 : tech / methodology / team size (white)
    _, cell_et = _merged_row(6, WHITE)
    p_tech = cell_et.paragraphs[0]
    p_tech.paragraph_format.space_before = Pt(3)
    p_tech.paragraph_format.space_after  = Pt(1)
    p_tech.paragraph_format.left_indent  = Cm(0.3)
    _run(p_tech, "Technologies : ", bold=True, size=9, color=TEAL_RGB)
    _run(p_tech, ", ".join(exp.technologies) if exp.technologies else "",
         size=9, color=BLACK_RGB)

    p_meth = cell_et.add_paragraph()
    p_meth.paragraph_format.space_before = Pt(1)
    p_meth.paragraph_format.space_after  = Pt(1)
    p_meth.paragraph_format.left_indent  = Cm(0.3)
    _run(p_meth, "Méthodologies : ", bold=True, size=9, color=TEAL_RGB)
    _run(p_meth, ", ".join(exp.methodologies) if exp.methodologies else "",
         size=9, color=BLACK_RGB)

    p_team = cell_et.add_paragraph()
    p_team.paragraph_format.space_before = Pt(1)
    p_team.paragraph_format.space_after  = Pt(3)
    p_team.paragraph_format.left_indent  = Cm(0.3)
    _run(p_team, "Taille de l'équipe : ", bold=True, size=9, color=TEAL_RGB)
    _run(p_team, str(exp.team_size) if exp.team_size else "",
         size=9, color=BLACK_RGB)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)


# ─────────────────────────────────────────────────────────────
#  GÉNÉRATION PRINCIPALE
# ─────────────────────────────────────────────────────────────

def generate_dossier_docx(dossier: DossierCompetences, output_dir: Path) -> Path:
    """Génère le dossier de compétences .docx au format template ACHMITECH."""

    doc = Document()

    # ── Marges ────────────────────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin   = Cm(2.0)
        sec.right_margin  = Cm(2.0)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    candidate       = dossier.candidate_name or "Candidat"
    extraction_date = (dossier.extraction_date or "")[:10]

    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    #  LIGNE 1 — Confidentiel ACHMITECH  |  ACHMITECH logo
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    top = doc.add_table(rows=1, cols=2)
    _no_borders(top)
    _table_width(top, PAGE_W)
    _col_width(top.rows[0].cells[0], PAGE_W * 0.6)
    _col_width(top.rows[0].cells[1], PAGE_W * 0.4)

    p_conf = top.rows[0].cells[0].paragraphs[0]
    p_conf.paragraph_format.space_after = Pt(2)
    _run(p_conf, "Confidentiel ACHMITECH", italic=True, size=8, color=GREY_RGB)

    p_logo = top.rows[0].cells[1].paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_logo.paragraph_format.space_after = Pt(2)
    _run(p_logo, "ACHMITECH", bold=True, size=11, color=TEAL_RGB)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    #  LIGNE 2 — Photo candidate  |  Titre
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    photo_row = doc.add_table(rows=1, cols=2)
    _no_borders(photo_row)
    _table_width(photo_row, PAGE_W)
    photo_cell = photo_row.rows[0].cells[0]
    title_cell = photo_row.rows[0].cells[1]
    _col_width(photo_cell, 3.5)
    _col_width(title_cell, PAGE_W - 3.5)

    # Photo
    photo_ok = False
    if dossier.photo and dossier.photo.found and dossier.photo.file_path:
        pp = Path(dossier.photo.file_path)
        if pp.exists():
            try:
                p_img = photo_cell.paragraphs[0]
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(4)
                p_img.add_run().add_picture(str(pp), width=Cm(2.8))
                photo_ok = True
            except Exception as exc:
                logger.warning("Photo non insérée: %s", exc)

    if not photo_ok:
        _cell_bg(photo_cell, "D0EAF8")
        p_ph = photo_cell.paragraphs[0]
        p_ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ph.paragraph_format.space_before = Pt(18)
        p_ph.paragraph_format.space_after  = Pt(18)
        _run(p_ph, "[ Photo ]", italic=True, size=9,
             color=RGBColor(0x55, 0x55, 0x55))

    # Titre
    p_title = title_cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(16)
    p_title.paragraph_format.space_after  = Pt(4)
    _run(p_title, "DOSSIER DE COMPETENCES",
         bold=True, size=18, color=TITLE_RGB)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    #  TABLEAU INFORMATIONS CANDIDAT
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    last_deg = "—"
    if dossier.last_degree:
        ld = dossier.last_degree
        last_deg = ld.degree + (f" ({ld.level})" if ld.level else "")

    INFO_LABEL = 4.5
    INFO_VALUE = PAGE_W - INFO_LABEL

    info_rows = [
        ("Nom Prénom",            candidate),
        ("Dernier Diplôme",       last_deg),
        ("Années d'Expériences",  _years_label(dossier)),
        ("Date de Disponibilité", ""),
    ]

    info_tbl = doc.add_table(rows=len(info_rows), cols=2)
    _no_borders(info_tbl)
    _table_width(info_tbl, PAGE_W)

    for i, (label, value) in enumerate(info_rows):
        lc = info_tbl.rows[i].cells[0]
        vc = info_tbl.rows[i].cells[1]
        _col_width(lc, INFO_LABEL)
        _col_width(vc, INFO_VALUE)
        _cell_bg(lc, TEAL)
        _cell_bg(vc, WHITE)
        _cell_borders(lc, "007777")
        _cell_borders(vc, BORDER_COLOR)

        p_l = lc.paragraphs[0]
        p_l.paragraph_format.space_before = Pt(3)
        p_l.paragraph_format.space_after  = Pt(3)
        p_l.paragraph_format.left_indent  = Cm(0.2)
        _run(p_l, label, bold=True, size=9, color=WHITE_RGB)

        p_v = vc.paragraphs[0]
        p_v.paragraph_format.space_before = Pt(3)
        p_v.paragraph_format.space_after  = Pt(3)
        p_v.paragraph_format.left_indent  = Cm(0.3)
        if value:
            _run(p_v, value, size=9, color=BLACK_RGB)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    #  VOS TOP 5 DES HARD SKILLS
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    _top5_table(doc, "VOS TOP 5 DES HARD SKILLS", [
        (sk.name, _stars(sk.level))
        for sk in (dossier.hard_skills or [])[:5]
    ])

    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    #  VOS TOP 5 DES SOFT SKILLS
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    _top5_table(doc, "VOS TOP 5 DES SOFT SKILLS", [
        (sk.name, _stars(sk.level))
        for sk in (dossier.soft_skills or [])[:5]
    ])

    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    #  VOS TOP 5 DES OUTILS MAÎTRISÉS
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    _top5_table(doc, "VOS TOP 5 DES OUTILS MAÎTRISÉS", [
        (t.name, _stars(t.level))
        for t in (dossier.top_tools or [])[:5]
    ])

    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    #  VOS LANGUES MAÎTRISÉES
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    lang_rows: list[tuple[str, str]] = []
    for lang in (dossier.languages or [])[:5]:
        if isinstance(lang, str):
            lang_rows.append((lang, ""))
        else:
            name  = getattr(lang, "name", str(lang))
            level = getattr(lang, "level", None)
            lbl   = getattr(lang, "level_label", None)
            level_str = lbl or (_stars(int(level)) if level else "")
            lang_rows.append((name, level_str))

    _top5_table(doc, "VOS LANGUES MAÎTRISÉES", lang_rows)

    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    #  PAGE 2  —  FORMATIONS
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    if dossier.educations:
        _add_page_break(doc)
        _section_banner(doc, "FORMATIONS")

        COL_W = [2.0, 9.5, PAGE_W - 11.5]
        edu_tbl = doc.add_table(rows=1, cols=3)
        _no_borders(edu_tbl)
        _table_width(edu_tbl, PAGE_W)

        for i, h in enumerate(["Année", "Diplôme", "Établissement"]):
            c = edu_tbl.rows[0].cells[i]
            _col_width(c, COL_W[i])
            _cell_bg(c, TEAL)
            _cell_borders(c, "007777")
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            _run(p, h, bold=True, size=9, color=WHITE_RGB)

        for idx, edu in enumerate(dossier.educations):
            row = edu_tbl.add_row()
            bg  = LIGHT_ROW if idx % 2 == 0 else WHITE
            vals = [
                str(edu.year) if edu.year else "—",
                (edu.degree or "—") + (" (en cours)" if edu.status == "en_cours" else ""),
                edu.school or "—",
            ]
            aligns = [WD_ALIGN_PARAGRAPH.CENTER,
                      WD_ALIGN_PARAGRAPH.LEFT,
                      WD_ALIGN_PARAGRAPH.LEFT]
            for ci, (txt, al) in enumerate(zip(vals, aligns)):
                c = row.cells[ci]
                _col_width(c, COL_W[ci])
                _cell_bg(c, bg)
                _cell_borders(c, BORDER_COLOR)
                p = c.paragraphs[0]
                p.alignment = al
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                p.paragraph_format.left_indent  = Cm(0.15)
                _run(p, txt, size=9, color=BLACK_RGB)

        doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    #  PAGE 2  —  EXPÉRIENCES PROFESSIONNELLES
    # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
    if dossier.experiences:
        _section_banner(
            doc,
            "EXPERIENCES (De la plus récente a la plus ancienne)"
        )
        for exp in dossier.experiences:
            _add_experience_block(doc, exp)

    # ── Pied de page ─────────────────────────────────────────
    footer_p = doc.sections[0].footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(footer_p,
         f"Cv2DC · Dossier de Compétences généré automatiquement · {extraction_date}",
         italic=True, size=8, color=GREY_RGB)

    # ── Sauvegarde ────────────────────────────────────────────
    safe = (dossier.candidate_name or "candidat").replace(" ", "_").lower()
    out  = output_dir / f"dossier_competences_{safe}.docx"
    doc.save(str(out))
    logger.info("✅ DOCX généré : %s", out)
    return out
